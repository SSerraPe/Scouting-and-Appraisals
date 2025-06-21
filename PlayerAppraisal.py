# -*- coding: utf-8 -*-
"""
Created on Fri Dec 22 09:45:32 2023

@author: Sebastián Serra
"""

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import networkx as nx
import numpy as np
from scipy.stats import percentileofscore

from docx import Document
from docx.shared import Inches
import io
import base64

def load_data(file_path):
    """
    Load a DataFrame from a CSV or Parquet file.

    Parameters:
    - file_path (str): The path to the input file.

    Returns:
    - pd.DataFrame: The loaded DataFrame.
    """
    if file_path.endswith('.csv'):
        return pd.read_csv(file_path)
    elif file_path.endswith('.parquet'):
        return pd.read_parquet(file_path)
    else:
        raise ValueError("Unsupported file format. Only CSV and Parquet files are supported.")

def plot_ridgeline(data, player_name, metrics):
    for metric in metrics:
        # Get the player's value for the specified metric
        player_value = data[data['Name'] == player_name][metric].iloc[0]

        # Calculate the percentile rank of the player's value
        percentile = np.percentile(data[metric], player_value)

        # Set up the figure and axes
        fig, ax = plt.subplots(figsize=(8, 6))

        # Define custom colors
        below_color = 'blue'
        above_color = 'green'

        # Plot the ridgeline plot
        sns.kdeplot(data=data[metric], color='blue', fill=True, ax=ax)

        # Change color of ridges based on player_value
        for line in ax.collections:
            vertices = line.get_paths()[0].vertices
            max_y = np.max(vertices[:, 1])
            if max_y > player_value:
                line.set_color(above_color)
            else:
                line.set_color(below_color)

        # Plot the red dashed line on the x-axis at the player's percentile
        plt.axvline(x=percentile, color='red', linestyle='dashed', linewidth=2)

        # Set labels and title
        plt.xlabel('Density')
        plt.ylabel(metric)
        plt.title('Ridgeline Plot for {}'.format(metric))

        plt.show()

def visualize_distribution(data, player_name, metric):
    """
    Visualize the distribution of a metric and mark the player's value with a vertical line.

    Parameters:
    - data (pd.DataFrame): The input DataFrame.
    - player_name (str): The player's name.
    - metric (str): The metric to visualize.
    """
    plt.figure(figsize=(10, 6))
    
    # Use a gradient background
    ax = plt.gca()
    ax.set_facecolor('#f5f5f5')  # light gray
    
    # Adjust the color palette
    sns.set_palette("pastel")
    
    # Plot the distribution with grid lines
    sns.histplot(data[metric], kde=True, bins=10, alpha=0.7)
    plt.grid(True, linestyle='--', alpha=0.5)
    
    # Get the player's value for the specified metric
    player_value = data[data['Name'] == player_name][metric].iloc[0]
    
    player_percentile = percentileofscore(data[metric], player_value)
    
    # Mark the player's value with a vertical line with a shadow effect
    plt.axvline(player_value, color='red', linestyle='dashed', linewidth=2, alpha=0.8, label=f'{player_name} Value', zorder=0)
    plt.axvline(player_value, color='black', linewidth=3, alpha=0.2, zorder=-1)  # shadow
    
    plt.title(f'Distribution of {metric} with {player_name} Percentile', fontsize=16)
    plt.xlabel(metric.capitalize(), fontsize=12)
    plt.ylabel('Frequency', fontsize=12)
    plt.legend()
    plt.tight_layout()  # Adjust layout to prevent clipping of titles and labels
    plt.show()
    
    return player_percentile

def vis_density(data, player_name, metrics):
    """
    Visualize the distribution of multiple metrics using half-violin plots and mark the player's value.

    Parameters:
    - data (pd.DataFrame): The input DataFrame.
    - player_name (str): The player's name.
    - metrics (list of str): The list of metrics to visualize.
    """
    plt.figure(figsize=(12, 8))
    pps = []
    
    for idx, metric in enumerate(metrics):
        plt.subplot(len(metrics), 1, idx + 1)
        
        # Get the player's value for the specified metric
        player_value = data[data['Name'] == player_name][metric].iloc[0]
        player_percentile = percentileofscore(data[metric], player_value)
        pps.append(player_percentile)
        
        # Plot individual half-violin plot with custom colors
        sns.violinplot(x=data[metric], inner='quartile', linewidth=1.5, cut=0, scale='width', bw=0.2, ax=plt.gca())
        
        # Mark the player's value with a vertical line
        plt.axvline(player_value, color='red', linestyle='dashed', linewidth=2)
        
        plt.xlabel('')
        plt.ylabel(metric, rotation=0, ha='right')

    plt.tight_layout()
    plt.show()
    
    return pps


def plot_radar(df, player_name, season, selected_metrics):
    """
    Generate a radar (spider) plot to visually represent a player's performance across selected metrics for a specific season.

    Parameters:
        - player_name (str): The name of the player for whom the radar plot is generated.
        - season (str): The season for which the radar plot is created.
        - selected_metrics (list): A list of metrics to include in the radar plot.
    """
    
    # Filter DataFrame based on player_name and season
    player_data = df[(df['Name'] == player_name) & (df['Season'] == season)]

    # Extract metric values for the selected metrics
    values = player_data[selected_metrics].values.flatten().tolist()

    # Number of selected metrics
    num_metrics = len(selected_metrics)

    # Calculate angle for each axis
    angles = np.linspace(0, 2 * np.pi, num_metrics, endpoint=False).tolist()

    # Close the plot (connect the first and last points)
    values += values[:1]
    angles += angles[:1]

    # Plot radar chart
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color='b', alpha=0.25)

    # Set labels for each axis
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(selected_metrics)

    # Set title
    ax.set_title(f'Radar Plot for {player_name} - {season}')

    # Show the plot
    plt.show()

def analyze_graph(graph):
    """
    Analyze a graph using density, average degree, degree centrality, and betweenness centrality.

    Parameters:
    - graph (networkx.Graph): The input graph.

    Returns:
    - analysis_results (dict): A dictionary containing analysis results.
    """

    # Density
    density = nx.density(graph)

    # Average Degree
    avg_degree = sum(dict(graph.degree()).values()) / len(graph)

    # Degree Centrality
    degree_centrality = nx.degree_centrality(graph)

    # Betweenness Centrality
    betweenness_centrality = nx.betweenness_centrality(graph)

    analysis_results = {
        'Density': density,
        'Average Degree': avg_degree,
        'Degree Centrality': degree_centrality,
        'Betweenness Centrality': betweenness_centrality
    }

    return analysis_results

def create_and_analyze_percentile_graph(data, player_name, season, metric_list):
    """
    Create and analyze a percentile graph for a specific player and season.

    Parameters:
    - data (pd.DataFrame): The input DataFrame containing metric data for all players.
    - player_name (str): The name of the player.
    - season (str): The specific season.
    - metric_list (list): A list of metrics to consider.
    - percentile_threshold (float): The percentile threshold for connecting nodes with an edge. Default is 0.75.

    Returns:
    - G (networkx.Graph): The constructed graph.
    - analysis_results (dict): Analysis results for the graph.
    """
    percentile_threshold=0.95
    
    G = nx.Graph()

    for metric in metric_list:
        player_value = data[data['Name'] == player_name][metric].iloc[0]
        player_percentile = percentileofscore(data[metric], player_value)
        G.add_node(metric, size=player_percentile)

    for metric1 in metric_list:
        for metric2 in metric_list:
            if metric1 != metric2:
                percentile1_v = data[data['Name'] == player_name][metric1].iloc[0]
                player_percentile1 = percentileofscore(data[metric], percentile1_v)
                percentile2_v = data[data['Name'] == player_name][metric2].iloc[0]
                player_percentile2 = percentileofscore(data[metric], percentile2_v)
                if max(player_percentile1, player_percentile2) > percentile_threshold:
                    G.add_edge(metric1, metric2)

    # Visualization
    pos = nx.circular_layout(G)
    node_sizes = [G.nodes[metric]['size'] * 10 for metric in G.nodes]
    nx.draw(G, pos, with_labels=True, font_weight='bold', node_size=node_sizes, node_color='skyblue', font_size=10, font_color='black', edge_color='gray', width=2)
    plt.title(f"Percentile Graph for {player_name} - {season}")
    plt.show()

    # Analysis
    analysis_results = analyze_graph(G)
    print("\nAnalysis Results:")
    for metric, value in analysis_results.items():
        print(f"{metric}: {value}")

    return G, analysis_results

"""
def create_results_document(file_path, player_name, season, selected_metrics):
    # Load data
    data = load_data(file_path)

    # Create a Word document
    doc = Document()

    # Visualize distribution for each selected metric
    for metric in selected_metrics:
        plt.figure(figsize=(10, 6))
        sns.histplot(data[metric], kde=True)
        player_percentile_value = data[data['Name'] == player_name][metric].rank(pct=True).iloc[0]
        plt.axvline(player_percentile_value, color='red', linestyle='dashed', linewidth=2, label=f'{player_name} Percentile')
        plt.title(f'Distribution of {metric} with {player_name} Percentile')
        plt.legend()
        image_stream = io.BytesIO()
        plt.savefig(image_stream, format='png')
        plt.close()

        # Add image to the Word document
        image_stream.seek(0)
        doc.add_picture(image_stream, width=Inches(5))
        doc.add_paragraph(f'\n{metric} Distribution')

    # Plot radar chart
    plt.figure(figsize=(10, 6))
    plot_radar(data, player_name, season, selected_metrics)
    image_stream = io.BytesIO()
    plt.savefig(image_stream, format='png')
    plt.close()

    # Add radar chart image to the Word document
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(5))
    doc.add_paragraph(f'\nRadar Chart for {player_name} - {season}')

    # Create and analyze percentile graph
    graph, analysis_results = create_and_analyze_percentile_graph(data, player_name, season, selected_metrics)

    # Add analysis results to the Word document
    doc.add_paragraph("\nAnalysis Results:")
    for metric, value in analysis_results.items():
        doc.add_paragraph(f"{metric}: {value}")

    # Save the document
    doc.save("results_document.docx")
"""


# Usage:

file_path = r"D:\AZTEC DH  COLLABORATIVE\Scouting & Appraisal\Data Repository\MLSNP_CAM.csv"
doc_path = "D:\AZTEC DH  COLLABORATIVE\Scouting & Appraisal"
data = load_data(file_path)

player_name = 'Jesús Armando Castellano Anuel'
season = 2023

for column in data.columns:
    print(column)

selected_metrics = [
    "Open Play xG Assisted",
    "xG Assisted",
    "Fouls Won",
    "Successful Dribbles",
    "Shots",
    "xG",
    "Successful Box Cross%",
    "Touches In Box",
    "xGBuildup"
]


pps = vis_density(data, player_name, selected_metrics)
print(pps)

pps = []
for metric in selected_metrics:
    pp = visualize_distribution(data, player_name, metric)
    pps.append(pp)
    
print(pps)

plot_radar(data, player_name, season, selected_metrics)

graph, results = create_and_analyze_percentile_graph(data, player_name, season, selected_metrics)

# create_results_document(doc_path, player_name, season, selected_metrics)




